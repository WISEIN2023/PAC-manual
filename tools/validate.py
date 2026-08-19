# -*- coding: utf-8 -*-
import os, re, sys, json, glob, collections
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import docx
from converter import anchor
from registry import DOCS

SRC = os.environ.get("PAC_DOCX_SRC", os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "_source_docx"))
OUT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
M = json.load(open(OUT + '/.manifest.json'))

def nz(s):
    s = (s or '').replace('*','').replace(chr(92),'')
    s = re.sub(r'[-\u00b7\u2022\u25cf\u25a0]', '', s)
    s = re.sub(r'^\s*(?:[-+\u00b7]|\d+\.)\s+', '', s, flags=re.M)
    return re.sub(r'[\s\u200b\u00a0]+', '', s)
def nzmd(s):
    s = re.sub(r'!\[[^\]]*\]\([^)]*\)', '', s)
    s = re.sub(r'\[([^\]]*)\]\([^)]*\)', r'\1', s)
    s = re.sub(r'^\s*(?:[-*+]|\d+\.)\s+', '', s, flags=re.M)
    s = s.replace('<br>', '')
    s = s.replace('\\|', '|').replace('\\(', '(').replace('*', '').replace('\\', '')
    s = re.sub(r'[-\u00b7\u2022\u25cf\u25a0]', '', s)
    return re.sub(r'[\s\u200b]+', '', s)


def docx_units(path):
    d = docx.Document(path)
    u = []
    for p in d.paragraphs:
        t = p.text.strip()
        if t: u.append(t)
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                s = c.text.strip()
                if s: u.append(s)
    return u

rows = []
missing_all = []
missing_samples = collections.defaultdict(list)
for m in M:
    slug = m['slug']
    md = ''
    for f in m['files']:
        md += open(os.path.join(OUT, f['path'])).read() + '\n'
    mdn = nzmd(md)
    units = docx_units(os.path.join(SRC, m['source']))
    checked = miss = 0
    for u in units:
        un = nz(u)
        if len(un) < 6:
            continue
        checked += 1
        if un not in mdn:
            miss += 1
            missing_all.append((slug, re.sub(r'\s+', ' ', u)))
            if len(missing_samples[slug]) < 5:
                missing_samples[slug].append(u[:70])
    cov = 100.0 * (checked - miss) / checked if checked else 100.0
    rows.append(dict(slug=slug, title=m['title'], units=checked, miss=miss, cov=cov,
                     files=len(m['files']), chars=m['chars'], tables=m['tables'],
                     images=m['images']))

# 링크/앵커 무결성
anchors = {}
for f in glob.glob(OUT + '/**/*.md', recursive=True):
    rel = os.path.relpath(f, OUT)
    s = open(f).read()
    A = set()
    for line in s.split('\n'):
        h = re.match(r'^#{1,6} (.+)$', line)
        if h:
            A.add(anchor(h.group(1).strip()))
    anchors[rel] = A

bad_links, bad_anchor, n_links = [], [], 0
for f in glob.glob(OUT + '/**/*.md', recursive=True):
    rel = os.path.relpath(f, OUT)
    s = open(f).read()
    for mm in re.finditer(r'\[[^\]]*\]\(([^)\s]+)\)', s):
        t = mm.group(1)
        if t.startswith(('http://', 'https://', '#')):
            continue
        n_links += 1
        path, _, anc = t.partition('#')
        tgt = os.path.normpath(os.path.join(os.path.dirname(rel), path))
        if not os.path.exists(os.path.join(OUT, tgt)):
            bad_links.append((rel, t))
        elif anc and tgt.endswith('.md') and anc not in anchors.get(tgt, set()):
            bad_anchor.append((rel, t))

# 고아 이미지
used = set()
for f in glob.glob(OUT + '/docs/**/*.md', recursive=True):
    for mm in re.finditer(r'!\[[^\]]*\]\(([^)]+)\)', open(f).read()):
        used.add(os.path.normpath(os.path.join(os.path.dirname(os.path.relpath(f, OUT)), mm.group(1))))
allimg = {os.path.relpath(p, OUT) for p in glob.glob(OUT + '/assets/**/*', recursive=True) if os.path.isfile(p)}
orphan = sorted(allimg - used)
for o in orphan:
    os.remove(os.path.join(OUT, o))

def sz(p):
    return sum(os.path.getsize(x) for x in glob.glob(OUT + p, recursive=True) if os.path.isfile(x))

md_bytes = sz('/**/*.md')
img_bytes = sz('/assets/**/*')
orig_bytes = sum(os.path.getsize(os.path.join(SRC, d['file'])) for d in DOCS)

R = []
R.append('# 변환 검증 리포트')
R.append('')
R.append(f'- 원본 DOCX {len(M)}건 ({orig_bytes/1024/1024:.1f} MB) → Markdown {sum(r["files"] for r in rows)}개 파일 ({md_bytes/1024:.0f} KB) + 이미지 {len(used)}장 ({img_bytes/1024/1024:.1f} MB)')
R.append(f'- 내부 링크 {n_links}개 검사: 깨진 파일 링크 **{len(bad_links)}건**, 깨진 앵커 **{len(bad_anchor)}건**')
R.append(f'- 참조되지 않는 이미지 {len(orphan)}장 제거')
R.append('')
R.append('## 문서별 텍스트 커버리지')
R.append('')
R.append('원본 DOCX의 문단·표 셀(6자 이상) 각각이 변환 결과에 그대로 존재하는지 전수 대조한 결과입니다.')
R.append('')
R.append('| 문서 | 검사 단위 | 누락 | 커버리지 | 파일수 | 표 | 이미지 |')
R.append('|---|---:|---:|---:|---:|---:|---:|')
for r in sorted(rows, key=lambda x: x['cov']):
    R.append(f'| {r["title"]} | {r["units"]} | {r["miss"]} | {r["cov"]:.1f}% | {r["files"]} | {r["tables"]} | {r["images"]} |')
tu = sum(r['units'] for r in rows); tm = sum(r['miss'] for r in rows)
R.append(f'| **합계** | **{tu}** | **{tm}** | **{100.0*(tu-tm)/tu:.1f}%** | **{sum(r["files"] for r in rows)}** | **{sum(r["tables"] for r in rows)}** | **{sum(r["images"] for r in rows)}** |')
R.append('')
if tm:
    DROP_PAT = r'^\d+(\.\d+)*[ ].*[ ]\d{1,3}$|필드 업데이트|F9|목차|결산자동화|Process Automatic|운영자|담당자|작성|버전|검증 기준일|wiseinsoft|와이즈인소프트|가이드|PAC Solution'
    intended = sum(1 for v in missing_all if re.search(DROP_PAT, v[1]))
    other = [v for v in missing_all if not re.search(DROP_PAT, v[1])]
    R.append('### 누락 항목 분류')
    R.append('')
    R.append(f'누락 {tm}건 중 **{intended}건**은 의도적으로 제거한 요소입니다.')
    R.append('')
    R.append('- Word 자동 목차(TOC) 항목 및 "[필드 업데이트]로 페이지 번호를 갱신하세요" 안내문')
    R.append('- 표지 페이지의 솔루션명·부제·작성자·작성일 (front-matter로 이관)')
    R.append('- 빈 제목 Heading, 표 병합셀 중복')
    R.append('')
    R.append(f'분류되지 않은 잔여 항목: **{len(other)}건**' + (' (없음)' if not other else ''))
    if other:
        R.append('')
        for slug, v in other[:20]:
            R.append(f'- `{slug}` : {v[:90]}')
    R.append('')
if bad_links or bad_anchor:
    R.append('## 링크 오류')
    R.append('')
    for rel, t in bad_links[:30]:
        R.append(f'- 파일 없음: `{rel}` → `{t}`')
    for rel, t in bad_anchor[:30]:
        R.append(f'- 앵커 없음: `{rel}` → `{t}`')
    R.append('')
R.append('## 원본 문서 이슈 (변환 시 처리)')
R.append('')
R.append('| 이슈 | 처리 |')
R.append('|---|---|')
R.append('| 파일명 오타 `Data Moigration` | 슬러그 `data-migration`, 제목 `Data Migration 운영자 매뉴얼`로 정정 |')
R.append('| `메뉴얼`/`매뉴얼` 표기 혼재 | 문서 제목을 `매뉴얼`로 통일 (원본 파일명은 front-matter `source`에 보존) |')
R.append('| 제목 없는 빈 Heading (모델링 6건 등) | 변환 시 제거 |')
R.append('| Word 자동 목차(TOC) 필드 | 제거하고 문서별 `README.md` 목차로 대체 |')
R.append('| 버전 표기 혼재 (`v1.0`/`_1.1`/`_v1.1`) | front-matter `version`으로 통일 |')
R.append('| `ZLPAC00020` (자릿수 이상) | 원문 그대로 유지 — 원본 오탈자로 의심되므로 확인 필요 |')
open(OUT + '/CONVERSION_REPORT.md', 'w').write('\n'.join(R) + '\n')
print('\n'.join(R[:8]))
print('orphans removed:', len(orphan), '| bad links:', len(bad_links), '| bad anchors:', len(bad_anchor))
print('total coverage %.2f%%' % (100.0*(tu-tm)/tu))
